import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_report(image_path: str, analysis_json: dict, save_path: str):
    """
    Generates a professional DOCX report from the JSON analysis.
    This version is more robust at parsing the AI-generated markdown.

    Args:
        image_path (str): The file path to the original drawing image.
        analysis_json (dict): The JSON output from the pluto_workflow.
        save_path (str): The full file path where the .docx report will be saved.
    """
    document = Document()

    # --- Section 1: Title and Image ---
    title = document.add_heading('HTP "Person" Drawing Assessment Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Add a little space
    
    try:
        document.add_picture(image_path, width=Inches(5.5))
        # Center the image
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        document.add_paragraph(f"[Could not load image: {e}]")

    document.add_paragraph() # Add some space

    # --- Section 2: Assessment Header ---
    p = document.add_paragraph()
    p.add_run('Assessment:').bold = True
    p.add_run('\nObservation')

    # --- Section 3: Summary ---
    summary_heading = document.add_heading('Summary:', level=2)
    final_report_text = analysis_json.get('final', '')

    # Robustly find the summary text
    summary_text = "Summary not found in AI output." # Default message
    # Use regex to find the summary section more reliably
    summary_match = re.search(r'### Overall Summary\s*\n([\s\S]*)', final_report_text, re.IGNORECASE)
    if summary_match:
        summary_text = summary_match.group(1).strip()
    
    document.add_paragraph(summary_text)

    # --- Section 4: Detailed Analysis ---
    detailed_analysis_heading = document.add_heading('Detailed Analysis:', level=2)
    
    # Robustly find the main analysis section (everything before the summary)
    analysis_text = final_report_text
    summary_start_index = final_report_text.lower().find('### overall summary')
    if summary_start_index != -1:
        analysis_text = final_report_text[:summary_start_index]

    # Process each line of the analysis section
    lines = analysis_text.strip().split('\n')
    
    for line in lines:
        cleaned_line = line.strip()
        
        # Skip empty lines or the main title
        if not cleaned_line or "### HTP" in cleaned_line:
            continue
            
        # Check for Feature headings (e.g., **Feature: Figure Size and Placement**)
        if cleaned_line.startswith('**Feature:'):
            # Add as a main bullet point
            feature_name = cleaned_line.replace('**Feature:', '').replace('**', '').strip()
            p = document.add_paragraph(style='List Bullet')
            p.add_run('Feature: ' + feature_name).bold = True
            
        # Check for Observation or Interpretation lines
        elif cleaned_line.startswith('*   **Observation**:') or cleaned_line.startswith('*   **Interpretation**:'):
            # Add as an indented sub-bullet point
            # This is more complex to do perfectly in docx, so we'll simplify the formatting
            # to ensure content is always present.
            
            # Let's simplify and add as a regular paragraph with indentation
            content = cleaned_line.replace('*   ', '').replace('**', '').strip()
            p = document.add_paragraph(content)
            p.paragraph_format.left_indent = Inches(0.5)

        # Handle any other lines that might be part of a multi-line interpretation
        elif cleaned_line:
            # Add it to the previous paragraph if it seems to be a continuation
            if document.paragraphs[-1].text:
                 document.paragraphs[-1].add_run(' ' + cleaned_line)
            else:
                 document.add_paragraph(cleaned_line, style='List Bullet 2')


    # --- Section 5: Save the document ---
    try:
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
        print(f"Report successfully saved to: {save_path}")
    except Exception as e:
        print(f"Error saving document at '{save_path}': {e}")
        traceback.print_exc()